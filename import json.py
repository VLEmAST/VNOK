import json
import numpy as np
import pymorphy3
import re
import string
import fasttext
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
import flet as ft
from flet import (
    Page,
    TextField,
    ElevatedButton,
    Text,
    Column,
    Row,
    FilePicker,
    FilePickerResultEvent,
    TextButton,
    ListView,
    Colors,
    Dropdown,
    dropdown,
    SnackBar,  # Import SnackBar
)
import threading
import os
import time  # Import time module
import io  # Добавлен импорт io

try:  # Попытка импортировать python-docx, если не установлен, то пользователю будет предложено установить
    from docx import Document
    from docx.shared import Inches  # Import Inches for setting image size
except ImportError:
    print(
        "Библиотека python-docx не установлена.  Установите ее с помощью pip install python-docx для сохранения в формате .doc"
    )
    Document = None  # Устанавливаем в None, чтобы избежать ошибок при попытке сохранения в doc


# Инициализация
morph = pymorphy3.MorphAnalyzer()

# Глобальные переменные
stop_words = set()
all_descriptions = []
all_names = []
# all_ids = []  # Удалена переменная all_ids
model = None
df = None
punctuation_regex = re.compile(f"[{re.escape(string.punctuation)}]")
loaded_files = []  # Список для хранения путей к загруженным файлам

# Пути по умолчанию для сохранения модели и результатов
model_save_path = "fasttext_model.bin"
output_folder = ""  # Папка для сохранения результатов (изначально пустая)

# Глобальные переменные
selected_file_index = 0  # Индекс файла для анализа
reference_text = ""  # Текст "опорного" описания
selected_format = "Текстовый файл"  # Формат сохранения по умолчанию
page = None  # Глобальная переменная для Page
results_save_file_picker = None #FilePicker для сохранения результатов
model_save_file_picker = None #FilePicker для сохранения модели
general_output = None  # Объявляем general_output как глобальную переменную
get_data_path = None
get_stopwords_path = None
file_dropdown_changed = None
format_changed = None
pick_files_data = None
pick_files_stopwords = None
show_results_clicked = None
save_results_button_clicked = None
save_model_button_clicked = None
pick_files_model = None
get_model_path = None
train_model_threaded = None
results_list = []  # Список для хранения результатов
full_file_name_text = None  # Text element to display the full file name


# Функция предобработки текста
def preprocess_text(text, stop_words, morph, punctuation_regex):
    text = re.sub(r"[\r\n]+", " ", text)
    text = punctuation_regex.sub("", text).lower()
    words = text.split()
    lemmatized_words = [
        morph.parse(word)[0].normal_form for word in words if word not in stop_words
    ]
    return " ".join(lemmatized_words)


# Функция загрузки данных из JSON файла
def load_data(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)

        if isinstance(data, dict):
            # Если JSON - это отдельный объект, создаем список, создаем список, содержащий этот объект
            data = [data]  # Преобразуем в список

        if isinstance(data, list):
            return data
        else:
            print(
                f"Ошибка: Файл '{filepath}' имеет неправильную структуру. Ожидается список или объект."
            )
            return None

    except FileNotFoundError:
        print(f"Ошибка: Файл '{filepath}' не найден.")
        return None
    except json.JSONDecodeError:
        print(f"Ошибка: Некорректный формат JSON в файле '{filepath}'.")
        return None
    except Exception as e:
        print(f"Произошла ошибка при загрузке данных: {e}")
        return None


def load_data_from_files(filepaths):
    global all_descriptions, all_names, df, loaded_files
    all_descriptions = []
    all_names = []
    # all_ids = []  # Удалена переменная all_ids
    loaded_files = filepaths  # Сохраняем список файлов
    success = True

    for filepath in filepaths:
        data = load_data(filepath)
        if data:
            if isinstance(data, list):
                for item in data:
                    if (
                        isinstance(item, dict)
                        and "Description" in item
                        and "Name" in item
                    ):  # Удаляем проверку на Id
                        all_descriptions.append(item.get("Description", ""))
                        all_names.append(item.get("Name", ""))
                        # all_ids.append(item.get("Id", None))  # Удаляем добавление Id
                        # if 'Id' in item:
                        #     del item['Id'] # Удаляем Id из item
                    else:
                        print(
                            f"Ошибка: Файл '{filepath}' содержит объекты без ключей 'Description' и 'Name'."  # Удаляем упоминание Id
                        )
                        success = False
                        break
            else:
                print(
                    f"Ошибка: Файл '{filepath}' содержит данные в неправильном формате. Ожидается список объектов."
                )
                success = False
                break
        if not success:
            break

    if success and all_descriptions:
        preprocessed_names = [
            preprocess_text(name, stop_words, morph, punctuation_regex)
            for name in all_names
        ]
        preprocessed_descriptions = [
            preprocess_text(desc, stop_words, morph, punctuation_regex)
            for desc in all_descriptions
        ]

        df = pd.DataFrame(
            {
                # 'id': all_ids,  # Удаляем столбец id
                "name": preprocessed_names,
                "description": preprocessed_descriptions,
            }
        )
        return True
    else:
        return False


# Функция загрузки стоп-слов
def load_stopwords(filepath):
    global stop_words
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            loaded_stopwords = json.load(f)
        if isinstance(loaded_stopwords, list):
            stop_words = set(loaded_stopwords)
        elif (
            isinstance(loaded_stopwords, dict) and "stop_words" in loaded_stopwords
        ):
            stop_words = set(loaded_stopwords["stop_words"])
        else:
            print(
                "Предупреждение: Неверный формат stopwords.json. Используется пустой список стоп-слов."
            )
    except FileNotFoundError:
        print(f"Ошибка: Файл 'stopwords.json' не найден.")
        return None
    except json.JSONDecodeError:
        print(f"Ошибка: Некорректный формат JSON в файле '{filepath}'.")
        return None
    except Exception as e:
        print(f"Произошла ошибка при загрузке стоп-слов: {e}")
        return None


def find_most_similar(
    reference_text,
    analysis_file_index,
    loaded_files,
    model,
    stop_words,
    morph,
    punctuation_regex,
    top_n=3,
):
    """Находит top_n наиболее похожих описаний в analysis_file_index относительно reference_text."""
    try:
        target_desc = preprocess_text(
            reference_text, stop_words, morph, punctuation_regex
        )
        target_vector = model.get_sentence_vector(target_desc).reshape(1, -1)

        # Загружаем данные для анализа из выбранного файла
        analysis_filepath = loaded_files[analysis_file_index]
        with open(analysis_filepath, "r", encoding="utf-8") as f:
            analysis_data = json.load(f)

        if not isinstance(analysis_data, list):
            print(f"Ошибка: Файл '{analysis_filepath}' не содержит список объектов.")
            return ["Ошибка: Файл анализа имеет неправильный формат."]

        analysis_descriptions = [item.get("Description", "") for item in analysis_data]
        analysis_names = [item.get("Name", "") for item in analysis_data]
        # analysis_ids = [item.get("Id", None) for item in analysis_data] #Удаляем

        similarities = []
        for i, desc in enumerate(analysis_descriptions):
            processed_desc = preprocess_text(
                analysis_names[i] + " " + desc, stop_words, morph, punctuation_regex
            )
            vector = model.get_sentence_vector(processed_desc).reshape(1, -1)
            similarity = cosine_similarity(target_vector, vector)[0][0]
            similarities.append(similarity)

        similarities = np.array(similarities)
        most_similar_indices = np.argsort(similarities)[-top_n:][::-1]

        results = [
            f"Опорное описание: '{reference_text}'\n"
            f"Наиболее похоже на описание: '{analysis_descriptions[i]}' (Name: {analysis_names[i]})\n"  # Удаляем Id
            f"(Сходство: {similarities[i]:.4f})"
            for i in most_similar_indices
        ]
        return results
    except FileNotFoundError:
        print(f"Ошибка: Один из файлов не найден.")
        return ["Ошибка: Файл не найден."]
    except json.JSONDecodeError:
        print(f"Ошибка: Некорректный формат JSON в одном из файлов.")
        return ["Ошибка: Некорректный формат JSON."]
    except Exception as e:
        print(f"Произошла ошибка: {e}")
        return [f"Произошла ошибка: {e}"]


# Функция обучения FastText
def train_model():
    global model
    try:
        # Создаем или открываем файл train.txt для обучения модели
        train_file_path = "train.txt"  # Определяем путь к файлу
        # 'a' - append mode, добавляем данные в конец файла
        with open(train_file_path, "a", encoding="utf-8") as f:
            for name, desc in zip(df["name"], df["description"]):
                f.write(name + "\n")
                f.write(desc + "\n")

        # Обучаем модель
        # Если модель уже загружена, используем ее, иначе создаем новую
        if model is None:
            model = fasttext.train_unsupervised(
                train_file_path,
                model="skipgram",
                dim=200,
                epoch=20,
                ws=5,
                minCount=5,
            )
        else:
            # Дообучение - это сложная задача для fasttext.  В данном случае,
            # наиболее простой вариант - просто добавить новые данные в файл
            # и переобучить модель с нуля.  Однако, это может быть неэффективно
            # для больших объемов данных.

            # В fasttext нет явной функции для "дообучения".
            # Решением может быть только повторное обучение на всём датасете.

            model = fasttext.train_unsupervised(
                train_file_path,
                model="skipgram",
                dim=200,
                epoch=20,
                ws=5,
                minCount=5,
            )

        # Удаляем файл train.txt после его использования
        if os.path.exists(train_file_path):
            os.remove(
                train_file_path
            )  # Добавляем сообщение об успешном удалении
            print(f"Файл {train_file_path} успешно удален.")
        else:
            print(f"Файл {train_file_path} не существует.")  # Сообщаем, если файл не найден

        return True
    except Exception as e:
        print(f"Произошла ошибка при обучении модели: {e}")
        return False


def save_results_to_file(results_list, output_path, file_format):
    try:
        if file_format == "Текстовый файл":
            with open(output_path, "w", encoding="utf-8") as f:
                for line in results_list:
                    f.write(line + "\n\n")  # Добавляем разделение между результатами

        elif file_format == "Документ Word":
            if Document is None:  # Проверяем, что python-docx установлен
                print(
                    "Невозможно сохранить в формате Word.  Библиотека python-docx не установлена."
                )
                return False
            document = Document()
            for line in results_list:
                document.add_paragraph(line)
                document.add_paragraph("")  # Разделение между результатами

            document.save(output_path)
        return True
    except Exception as e:
        print(f"Произошла ошибка при сохранении результатов: {e}")
        return False


def load_model(path):
    global model
    try:
        print(f"Попытка загрузить модель из: {path}")  # Выводим путь
        if not os.path.exists(path):
            print(f"Ошибка: Файл не найден по пути {path}")
            return False
        model = fasttext.load_model(path)
        print("Модель успешно загружена.")
        return True
    except Exception as e:
        print(f"Произошла ошибка при загрузке модели: {e}")
        return False


def load_data_threaded(filepaths):  # Добавлена функция обратного вызова
    global is_loading_data, all_descriptions, all_names, df
    is_loading_data = True
    update_ui()

    success = load_data_from_files(filepaths)
    if success:
        general_output.value = f"Данные успешно загружены из файлов"
        show_snack_bar("Данные успешно загружены.")  # Show success SnackBar
    else:
        general_output.value = f"Ошибка при загрузке данных из файлов"

    is_loading_data = False
    update_ui()


def save_model():
    global model, model_save_file_picker
    try:
        if model is not None:
            file_extension = ".bin"
            model_save_file_picker.save_file(file_name=f"fasttext_model{file_extension}")  # Запускаем FilePicker для выбора пути
            return True
        else:
            print("Ошибка: Модель не загружена.")
            return False
    except Exception as e:
        print(f"Произошла ошибка при сохранении модели: {e}")
        return False

def save_results(): #Функция для запуска сохранения результатов
    global results_save_file_picker
    file_extension = (
            ".txt" if selected_format == "Текстовый файл" else ".doc"
        )  # Определяем расширение файла
    results_save_file_picker.save_file(file_name=f"results{file_extension}")

def get_results_save_path(e: FilePickerResultEvent):
    global output_folder, results_list, general_output  # Добавил results_list
    if e.path:
        output_path = e.path
        try:
            save_results_to_file(
                results_list, output_path, selected_format
            )  # Теперь передаем results_list
            general_output.value = f"Результаты успешно сохранены в {output_path}"
            show_snack_bar("Результаты успешно сохранены.")  # Show success SnackBar
        except Exception as e:
            general_output.value = f"Произошла ошибка при сохранении результатов: {e}"

    else:
        general_output.value = "Выбор пути сохранения отменен."
    update_ui()

def get_model_save_path(e: FilePickerResultEvent):
    global model, general_output
    if e.path:
        output_path = e.path
        try:
            model.save_model(output_path)  # Сохраняем модель
            general_output.value = f"Модель успешно сохранена в {output_path}"
            show_snack_bar("Модель успешно сохранена.")
        except Exception as e:
            general_output.value = f"Произошла ошибка при сохранении модели: {e}"
    else:
        general_output.value = "Выбор пути сохранения отменен."
    update_ui()


# Объявляем update_ui как глобальную функцию
update_ui = None


def main(page_: ft.Page):
    global page, update_ui, results_save_file_picker, model_save_file_picker, general_output, get_data_path, get_stopwords_path, file_dropdown_changed, format_changed, pick_files_data, pick_files_stopwords, show_results_clicked, save_results_button_clicked, save_model_button_clicked, pick_files_model, get_model_path, train_model_threaded, show_snack_bar, results_list, full_file_name_text, selected_format  # Используем глобальные переменные


    page = page_
    page.title = "Анализ текстовых описаний"
    page.vertical_alignment = ft.MainAxisAlignment.START
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER

    # Инициализация состояний
    output_text = ""
    is_loading_data = False
    is_training_model = False
    global selected_file_index, reference_text, model  # Объявляем global

    page.snack_bar = SnackBar(content=Text(""))  # Init SnackBar

    def update_ui():
        page.update()

    # Use page.show_snack() to display SnackBar
    def show_snack_bar(message):
        page.snack_bar.content = Text(message)
        page.snack_bar.open = True
        update_ui()

    def get_file_name(filepath):
        return os.path.basename(filepath)

    def truncate_string(text, max_length):
        if len(text) > max_length:
            return text[: max_length - 3] + "..."
        else:
            return text

    # UI elements
    def get_data_path(e: FilePickerResultEvent):
        global is_loading_data, loaded_files
        is_loading_data = True
        update_ui()

        if e.files:
            filepaths = [f.path for f in e.files]
            general_output.value = f"Выбрано файлов данных: {len(filepaths)}"
            load_data_threaded(filepaths)

            data_button.disabled = False
            data_button.bgcolor = None

            # После завершения загрузки данных обновляем выпадающие списки
            file_dropdown.options.clear()
            for i, filepath in enumerate(loaded_files):
                filename = get_file_name(filepath)  # Get only the file name
                truncated_filename = truncate_string(
                    filename, 25
                )  # Обрезаем имя файла
                file_dropdown.options.append(
                    dropdown.Option(key=i, text=truncated_filename)
                )  # Используем обрезанное имя
            file_dropdown.value = 0
            if loaded_files:
                full_file_name_text.value = loaded_files[0]
            else:
                full_file_name_text.value = ""
            show_snack_bar("Файлы данных успешно загружены.")  # Show success SnackBar
            is_loading_data = False
            update_ui()
        else:
            general_output.value = "Выбор файла данных отменен!"
            data_button.disabled = False
            data_button.bgcolor = None
            is_loading_data = False
            update_ui()

    def get_stopwords_path(e: FilePickerResultEvent):
        global is_loading_data
        is_loading_data = True
        update_ui()

        if e.files:
            filepath = e.files[0].path
            general_output.value = f"Выбран файл стоп-слов: {filepath}"
            # success = load_stopwords_threaded(filepath)
            success = load_stopwords(filepath)
            if not success:
                stopwords_button.disabled = False
                stopwords_button.bgcolor = None
            show_snack_bar("Файл стоп-слов успешно загружен.")  # Show success SnackBar
        else:
            general_output.value = "Выбор файла стоп-слов отменен!"
            stopwords_button.disabled = False
            stopwords_button.bgcolor = None
        is_loading_data = False
        update_ui()

    def file_dropdown_changed(e):
        global selected_file_index, full_file_name_text
        selected_file_index = int(file_dropdown.value)
        full_file_name_text.value = loaded_files[selected_file_index]  # Устанавливаем полное имя файла
        general_output.value = (
            f"Выбран файл для анализа: {loaded_files[selected_file_index]}"
        )
        update_ui()

    def format_changed(e):
        global selected_format
        selected_format = format_dropdown.value
        print(f"Выбран формат: {selected_format}")
        update_ui()

    def pick_files_data(e):
        data_button.disabled = True
        data_button.bgcolor = ft.colors.GREEN
        update_ui()
        data_file_picker.pick_files(allow_multiple=True, allowed_extensions=["json"])

    def pick_files_stopwords(e):
        stopwords_button.disabled = True
        stopwords_button.bgcolor = ft.colors.GREEN
        update_ui()
        stopwords_file_picker.pick_files(allowed_extensions=["json"])

    def show_results_clicked(e):
        global model, loaded_files, selected_file_index, reference_text, results_list

        reference_text = reference_text_field.value  # Получаем текст из поля

        if not reference_text:
            general_output.value = "Ошибка: Введите опорное описание."
            update_ui()
            return

        if model is None:
            general_output.value = "Ошибка: Сначала обучите модель."
            update_ui()
            return

        results_list = find_most_similar(
            reference_text,
            selected_file_index,
            loaded_files,
            model,
            stop_words,
            morph,
            punctuation_regex,
        )

        if not isinstance(results_list, list):  # Исправил тут
            general_output.value = results_list
            results_text.controls.clear()  # Очищаем предыдущие результаты
            show_snack_bar("Произошла ошибка во время анализа.")  # Show error SnackBar
        else:
            results_text.controls.clear()
            for res in results_list:
                results_text.controls.append(Text(value=res))
            general_output.value = "Результаты анализа получены."
            show_snack_bar("Результаты успешно получены.")  # Show success SnackBar

        update_ui()

    def save_results_button_clicked(e):
        save_results()

    def save_model_button_clicked(e):
        save_model()

    def pick_files_model(e):
        model_button.disabled = True
        model_button.bgcolor = ft.colors.GREEN
        update_ui()
        model_file_picker.pick_files(allowed_extensions=["bin"])

    def get_model_path(e: FilePickerResultEvent):
        if e.files:
            filepath = e.files[0].path
            general_output.value = f"Выбран файл модели: {filepath}"
            if load_model(filepath):
                general_output.value = f"Модель успешно загружена из файла {filepath}."
                show_snack_bar("Модель успешно загружена.")  # Show success SnackBar
            else:
                general_output.value = "Ошибка при загрузке модели."
                show_snack_bar("Произошла ошибка во время загрузки модели.")  # Show success SnackBar
        else:
            general_output.value = "Выбор файла модели отменен!"
        model_button.disabled = False
        model_button.bgcolor = None
        update_ui()

    def train_model_threaded():
        global is_training_model, model
        is_training_model = True
        train_button.disabled = True
        train_button.bgcolor = ft.colors.GREEN
        update_ui()

        success = train_model()

        train_button.disabled = False
        train_button.bgcolor = None
        if success:
            general_output.value = "Модель FastText успешно обучена."  # Сообщение об успехе
            show_snack_bar("Модель FastText успешно обучена.")  # Show success SnackBar
        else:
            general_output.value = "Ошибка при обучении модели."  # Сообщение об ошибке

        is_training_model = False
        update_ui()

    globals()["get_data_path"] = get_data_path
    globals()["get_stopwords_path"] = get_stopwords_path
    globals()["file_dropdown_changed"] = file_dropdown_changed
    globals()["format_changed"] = format_changed
    globals()["pick_files_data"] = pick_files_data
    globals()["pick_files_stopwords"] = pick_files_stopwords
    globals()["show_results_clicked"] = show_results_clicked
    globals()["save_results_button_clicked"] = save_results_button_clicked
    globals()["save_model_button_clicked"] = save_model_button_clicked
    globals()["pick_files_model"] = pick_files_model
    globals()["get_model_path"] = get_model_path
    globals()["train_model_threaded"] = train_model_threaded
    globals()["full_file_name_text"] = full_file_name_text  # Добавлено чтобы не было ошибки
    globals()["selected_format"] = selected_format  # Добавлено чтобы не было ошибки
    globals()["results_save_file_picker"] = results_save_file_picker
    globals()["model_save_file_picker"] = model_save_file_picker

    data_file_picker = FilePicker(on_result=get_data_path)
    stopwords_file_picker = FilePicker(on_result=get_stopwords_path)
    results_save_file_picker = FilePicker(on_result=get_results_save_path) #FilePicker для сохранения результатов
    model_save_file_picker = FilePicker(on_result=get_model_save_path) #FilePicker для сохранения модели
    model_file_picker = FilePicker(on_result=get_model_path)

    file_dropdown = Dropdown(
        options=[],
        label="Файл для анализа",
        on_change=file_dropdown_changed,
        value=None,
        width=300,
    )

    reference_text_field = TextField(
        label="Опорное описание", multiline=True, width=500
    )

    format_dropdown = Dropdown(
        options=[dropdown.Option("Текстовый файл"), dropdown.Option("Документ Word")],
        value="Текстовый файл",
        label="Формат сохранения",
        on_change=format_changed,
        width=200,
    )

    # Кнопки
    train_button = ElevatedButton(
        "Обучить модель FastText",
        on_click=lambda _: threading.Thread(target=train_model_threaded).start(),
        disabled=is_training_model,
    )
    data_button = ElevatedButton("Загрузить данные", on_click=pick_files_data)
    stopwords_button = ElevatedButton(
        "Загрузить стоп-слова (stopwords.json)", on_click=pick_files_stopwords
    )

    model_button = ElevatedButton("Загрузить модель", on_click=pick_files_model)
    show_results_button = ElevatedButton("Вывести результаты", on_click=show_results_clicked)
    save_results_button = ElevatedButton(
        "Сохранить результаты", on_click=save_results_button_clicked
    )  # Кнопка "Сохранить результаты"
    save_model_button = ElevatedButton(
        "Сохранить модель", on_click=save_model_button_clicked
    )  # Кнопка "Сохранить модель"

    # Инициализируем general_output перед использованием
    general_output = Text("")

    # Инициализируем results_text перед использованием
    results_text = ListView(expand=True, spacing=10, auto_scroll=True)

    # Text element to display the full file name
    full_file_name_text = Text(value="")

    page.overlay.append(data_file_picker)
    page.overlay.append(stopwords_file_picker)
    page.overlay.append(results_save_file_picker) #Добавлены FilePicker
    page.overlay.append(model_save_file_picker) #Добавлены FilePicker
    page.overlay.append(model_file_picker)

    page.add(
        ft.Row(
            [data_button, stopwords_button], alignment=ft.MainAxisAlignment.CENTER
        ),
        ft.Row([train_button, model_button], alignment=ft.MainAxisAlignment.CENTER),
        ft.Row(
            [show_results_button, save_results_button],
            alignment=ft.MainAxisAlignment.CENTER,
        ),
        ft.Row([save_model_button], alignment=ft.MainAxisAlignment.CENTER),
        reference_text_field,
        ft.Row(
            [file_dropdown, format_dropdown], alignment=ft.MainAxisAlignment.CENTER
        ),
        full_file_name_text,  # Added Text to display the full file name
        general_output,
        results_text,
    )
    # Присваиваем функцию update_ui глобальной переменной
    globals()["update_ui"] = update_ui

    # Init SnackBar
    page.snack_bar = SnackBar(content=Text(""))


if __name__ == "__main__":
    ft.app(target=main)
